from __future__ import annotations
from io import BytesIO
from pathlib import Path
import zipfile
import pandas as pd
from .models import DocumentRecord


def _safe_decode(data: bytes) -> str:
    for enc in ["utf-8", "utf-16", "latin-1"]:
        try:
            return data.decode(enc)
        except Exception:
            pass
    return data.decode("utf-8", errors="ignore")


def parse_pdf(data: bytes, filename: str) -> DocumentRecord:
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(data))
        texts = []
        for i, page in enumerate(reader.pages):
            try:
                texts.append(f"\n\n[Page {i+1}]\n" + (page.extract_text() or ""))
            except Exception:
                texts.append(f"\n\n[Page {i+1}]\n")
        text = "\n".join(texts).strip()
        return DocumentRecord(filename=filename, file_type="pdf", text=text, page_count=len(reader.pages), char_count=len(text))
    except Exception as exc:
        return DocumentRecord(filename=filename, file_type="pdf", status="error", error=str(exc))


def parse_docx(data: bytes, filename: str) -> DocumentRecord:
    try:
        from docx import Document
        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts).strip()
        return DocumentRecord(filename=filename, file_type="docx", text=text, char_count=len(text))
    except Exception as exc:
        return DocumentRecord(filename=filename, file_type="docx", status="error", error=str(exc))


def parse_table(data: bytes, filename: str) -> tuple[DocumentRecord, list[pd.DataFrame]]:
    try:
        if filename.lower().endswith(".csv"):
            df = pd.read_csv(BytesIO(data))
        else:
            df = pd.read_excel(BytesIO(data))
        text = f"Table with columns: {', '.join(map(str, df.columns))}\n\n" + df.head(50).to_csv(index=False)
        doc = DocumentRecord(filename=filename, file_type="table", text=text, char_count=len(text), table_count=1)
        return doc, [df]
    except Exception as exc:
        return DocumentRecord(filename=filename, file_type="table", status="error", error=str(exc)), []


def parse_bytes(data: bytes, filename: str) -> tuple[list[DocumentRecord], list[pd.DataFrame]]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return [parse_pdf(data, filename)], []
    if lower.endswith(".docx"):
        return [parse_docx(data, filename)], []
    if lower.endswith((".txt", ".md")):
        text = _safe_decode(data)
        return [DocumentRecord(filename=filename, file_type="text", text=text, char_count=len(text))], []
    if lower.endswith((".csv", ".xlsx", ".xls")):
        doc, tables = parse_table(data, filename)
        return [doc], tables
    if lower.endswith(".zip"):
        docs, tables = [], []
        try:
            with zipfile.ZipFile(BytesIO(data)) as z:
                for name in z.namelist():
                    if name.endswith("/") or name.startswith("__MACOSX/"):
                        continue
                    inner_docs, inner_tables = parse_bytes(z.read(name), Path(name).name)
                    docs.extend(inner_docs); tables.extend(inner_tables)
            return docs, tables
        except Exception as exc:
            return [DocumentRecord(filename=filename, file_type="zip", status="error", error=str(exc))], []
    return [DocumentRecord(filename=filename, file_type="unknown", status="skipped", error="Unsupported file type")], []


def load_sample_documents(base_dir: Path) -> list[DocumentRecord]:
    docs = []
    for p in (base_dir / "sample_data").glob("sample_*.txt"):
        text = p.read_text(encoding="utf-8")
        docs.append(DocumentRecord(filename=p.name, file_type="text", text=text, char_count=len(text), source="sample"))
    return docs


def load_rules(base_dir: Path) -> pd.DataFrame:
    return pd.read_csv(base_dir / "sample_data" / "validation_rules.csv")


def load_connectors(base_dir: Path) -> pd.DataFrame:
    return pd.read_csv(base_dir / "sample_data" / "connectors.csv").fillna("")
